import os
import re
from datetime import datetime
from typing import Dict, Optional

from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt
from openai import OpenAI

DEFAULT_OUTPUT_DIR = "generated_docs"
DEFAULT_TEMPLATE_PATH = "templates/cover_template.docx"
DEFAULT_MODEL = os.getenv("OPENAI_MODEL", "gpt-5-mini-2025-08-07")

client = OpenAI()


def _skills_highlight_by_resume_type(resume_type: str) -> str:
    resume_type = (resume_type or "").upper()
    if resume_type == "DATA":
        return (
            "Highlight strengths in data analysis, Python, SQL, dashboards, "
            "automation, statistical modeling, and applied ML."
        )
    if resume_type == "RESEARCH":
        return (
            "Highlight strengths in academic research, experimental design, "
            "technical writing, scientific programming, and interdisciplinary work."
        )
    if resume_type == "ASSOCIATE":
        return (
            "Highlight strengths in operations, coordination, project management, "
            "stakeholder communication, process improvement, and documentation."
        )
    return "Highlight the most relevant strengths for this role."


def generate_cover_letter_body_llm(
    job_data: Dict,
    resume_type: str,
    personal_info: Dict,
    tone_instructions: Optional[str] = None,
    template_prompt: Optional[str] = None,
    model: str = DEFAULT_MODEL,
) -> str:
    title = job_data.get("Job_Title", "the role")
    job_id = job_data.get("Job_ID", "")
    dept = job_data.get("Department", "the department")
    description = job_data.get("Description", "")

    skills_hint = _skills_highlight_by_resume_type(resume_type)

    tone_rules = tone_instructions or (
        "Professional, confident, warm, and specific. "
        "No generic filler, no obvious AI patterns. "
        "3–5 short paragraphs. "
        "Include 1–2 concrete achievements that match the job. "
        "Use first person. "
        "Do not include placeholders like [Company] or [Hiring Manager]. "
        "Do not repeat the header/address. "
        "End the body right before 'Sincerely' (do NOT write 'Sincerely')."
    )

    system_prompt = (
        "You are an expert career writer. "
        "Write a tailored cover-letter BODY that reads like a human wrote it carefully. "
        "Follow the user's tone and structure instructions exactly."
    )

    user_prompt = f"""
Write the cover letter body for this application.

Candidate:
- Name: {personal_info.get("first_name", "")} {personal_info.get("last_name", "")}
- Background focus: {resume_type}
- Guidance: {skills_hint}

Job:
- Title: {title}
- Job ID: {job_id}
- Department/Unit: {dept}
- Description:
{description}

Tone + structure instructions:
{tone_rules}

If a custom template/prompt is provided, follow it:
{template_prompt or "None"}
""".strip()

    resp = client.responses.create(
        model=model,
        input=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
        temperature=0.7,
        max_output_tokens=900,
    )

    return resp.output_text.strip()


def _ensure_dir(path: str):
    os.makedirs(path, exist_ok=True)


def _sanitize_filename(text: str) -> str:
    text = text or "Unknown"
    return re.sub(r"[^a-zA-Z0-9_-]+", "_", text).strip("_")


def _replace_placeholders(doc: Document, mapping: Dict[str, str]) -> bool:
    found_any = False

    def replace_in_paragraph(p):
        nonlocal found_any
        original = p.text
        new_text = original
        for k, v in mapping.items():
            if k in new_text:
                new_text = new_text.replace(k, v)
        if new_text != original:
            p.text = new_text
            found_any = True

    for p in doc.paragraphs:
        replace_in_paragraph(p)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_in_paragraph(p)

    return found_any


def _add_body_paragraphs(doc: Document, body_text: str):
    chunks = [c.strip() for c in body_text.split("\n")]
    bullet_buffer = []

    for line in chunks:
        if not line:
            if bullet_buffer:
                for b in bullet_buffer:
                    doc.add_paragraph(b, style="List Bullet")
                bullet_buffer = []
            continue

        if line.startswith("- "):
            bullet_buffer.append(line[2:].strip())
        else:
            if bullet_buffer:
                for b in bullet_buffer:
                    doc.add_paragraph(b, style="List Bullet")
                bullet_buffer = []
            doc.add_paragraph(line)

    if bullet_buffer:
        for b in bullet_buffer:
            doc.add_paragraph(b, style="List Bullet")


def generate_cover_letter_docx(
    job_data: Dict,
    resume_type: str,
    personal_info: Dict,
    output_dir: str = DEFAULT_OUTPUT_DIR,
    template_path: Optional[str] = DEFAULT_TEMPLATE_PATH,
    tone_instructions: Optional[str] = None,
    template_prompt: Optional[str] = None,
    use_llm: bool = True,
) -> str:
    _ensure_dir(output_dir)

    date_str = datetime.now().strftime("%m/%d/%Y")
    title = job_data.get("Job_Title", "Position")
    job_id = job_data.get("Job_ID", "")
    dept = job_data.get("Department", "Hiring Committee")
    campus_or_location = job_data.get("Campus", "")

    if use_llm:
        body_text = generate_cover_letter_body_llm(
            job_data=job_data,
            resume_type=resume_type,
            personal_info=personal_info,
            tone_instructions=tone_instructions,
            template_prompt=template_prompt,
        )
    else:
        body_text = (
            f"I am excited to apply for the {title} role in {dept}. "
            "My background aligns well with your needs, and I’m confident I can contribute immediately.\n\n"
            "In my recent roles, I delivered measurable impact through analysis, execution, and collaboration. "
            "I’m motivated by UB’s mission and the opportunity to support your team’s work.\n\n"
            "Thank you for your consideration. I look forward to discussing fit and next steps."
        )

    doc = None
    template_used = False

    if template_path and os.path.exists(template_path):
        doc = Document(template_path)
        template_used = True
    else:
        doc = Document()

    mapping = {
        "{{DATE}}": date_str,
        "{{DEPARTMENT}}": dept,
        "{{TITLE}}": title,
        "{{JOB_ID}}": job_id,
        "{{BODY}}": body_text,
        "{{FIRST_NAME}}": personal_info.get("first_name", ""),
        "{{LAST_NAME}}": personal_info.get("last_name", ""),
        "{{EMAIL}}": personal_info.get("email", ""),
        "{{PHONE}}": personal_info.get("phone", ""),
        "{{LOCATION}}": personal_info.get("location", ""),
        "{{LINKEDIN}}": personal_info.get("linkedin", ""),
        "{{GITHUB}}": personal_info.get("github", ""),
        "{{CAMPUS}}": campus_or_location,
    }

    placeholders_found = _replace_placeholders(doc, mapping) if template_used else False

    if not placeholders_found:
        doc = Document()

        full_name = f"{personal_info.get('first_name','')} {personal_info.get('last_name','')}".strip()
        location = personal_info.get("location", "Potsdam, NY")
        phone = personal_info.get("phone", "")
        email = personal_info.get("email", "")
        linkedin = personal_info.get("linkedin", "LinkedIn")
        github = personal_info.get("github", "GitHub")

        p = doc.add_paragraph(full_name)
        p.runs[0].bold = True

        doc.add_paragraph(f"{location} | {phone}".strip(" |"))
        doc.add_paragraph(email)
        doc.add_paragraph(f"{linkedin} | {github}".strip(" |"))
        doc.add_paragraph("")
        doc.add_paragraph(date_str)
        doc.add_paragraph("")
        doc.add_paragraph("Hiring Committee")
        doc.add_paragraph(dept)
        doc.add_paragraph("University at Buffalo")
        if campus_or_location:
            doc.add_paragraph(campus_or_location)
        doc.add_paragraph("")
        doc.add_paragraph("Dear Hiring Committee,")
        doc.add_paragraph("")
        _add_body_paragraphs(doc, body_text)
        doc.add_paragraph("")
        doc.add_paragraph("Sincerely,")
        doc.add_paragraph(full_name)

    for style in doc.styles:
        if style.type == 1:
            try:
                style.font.name = "Times New Roman"
                style.font.size = Pt(12)
            except Exception:
                pass

    safe_job_id = _sanitize_filename(job_id)
    filename = os.path.join(output_dir, f"Cover_Letter_{safe_job_id}.docx")
    doc.save(filename)
    return filename

def generate_cover_letter(
    job_data: Dict,
    resume_type: str,
    personal_info: Dict,
    output_dir: str = DEFAULT_OUTPUT_DIR,
    template_path: Optional[str] = DEFAULT_TEMPLATE_PATH,
    tone_instructions: Optional[str] = None,
    template_prompt: Optional[str] = None,
    use_llm: bool = True,
) -> str:
    """
    High-level helper used by main.py.
    Generates a DOCX cover letter and returns the file path.
    """
    return generate_cover_letter_docx(
        job_data=job_data,
        resume_type=resume_type,
        personal_info=personal_info,
        output_dir=output_dir,
        template_path=template_path,
        tone_instructions=tone_instructions,
        template_prompt=template_prompt,
        use_llm=use_llm,
    )


def save_cover_letter(letter_docx_path: str) -> str:
    return letter_docx_path
